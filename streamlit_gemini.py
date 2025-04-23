# import streamlit as st
# import requests
# import json
# import io
# import os
# import tempfile
# import base64
# from pathlib import Path

# # Document processing libraries
# import PyPDF2
# import docx
# import fitz  # PyMuPDF
# from docx import Document

# # Set page config
# st.set_page_config(
#     page_title="Legal Document Translator (English → Hindi)",
#     page_icon="🔄",
#     layout="wide"
# )

# # Application title and description
# st.title("Legal Document Translator")
# st.markdown("### English to Hindi Legal Document Translation")
# st.markdown("Upload a document (.pdf, .txt, or .docx) to translate legal content from English to Hindi.")

# # API endpoint
# TRANSLATION_API_URL = "https://d89e-2401-4900-8843-57cc-cdb6-af8-25c0-bb20.ngrok-free.app/translate"

# # Function to extract text from PDF
# def extract_text_from_pdf(file):
#     text = ""
#     try:
#         # Try using PyMuPDF first (better for complex PDFs)
#         pdf_document = fitz.open(stream=file.read(), filetype="pdf")
#         for page_num in range(len(pdf_document)):
#             page = pdf_document.load_page(page_num)
#             text += page.get_text()
#         pdf_document.close()
#     except Exception as e:
#         st.warning(f"PyMuPDF extraction failed, trying PyPDF2: {e}")
#         # Reset file pointer
#         file.seek(0)
#         # Fallback to PyPDF2
#         pdf_reader = PyPDF2.PdfReader(file)
#         for page_num in range(len(pdf_reader.pages)):
#             text += pdf_reader.pages[page_num].extract_text() or ""
    
#     return text

# # Function to extract text from DOCX
# def extract_text_from_docx(file):
#     doc = Document(file)
#     return " ".join([paragraph.text for paragraph in doc.paragraphs])

# # Function to translate text
# def translate_text(text):
#     if not text.strip():
#         return "No text content found to translate."
    
#     try:
#         response = requests.post(
#             TRANSLATION_API_URL,
#             json={"text": text},
#             headers={"Content-Type": "application/json"}
#         )
        
#         if response.status_code == 200:
#             return response.json().get("translation", "Translation not found in response")
#         else:
#             return f"Error: {response.status_code} - {response.text}"
#     except Exception as e:
#         return f"Translation API error: {str(e)}"

# # Function to save text as PDF
# def save_as_pdf(text):
#     from reportlab.lib.pagesizes import letter
#     from reportlab.pdfgen import canvas
#     from reportlab.pdfbase import pdfmetrics
#     from reportlab.pdfbase.ttfonts import TTFont
#     from io import BytesIO
    
#     # Register Hindi font
#     try:
#         pdfmetrics.registerFont(TTFont('NotoSans', 'NotoSans-Regular.ttf'))
#         font_name = 'NotoSans'
#     except:
#         # Fallback to default font if Noto Sans is not available
#         font_name = 'Helvetica'
    
#     buffer = BytesIO()
#     c = canvas.Canvas(buffer, pagesize=letter)
#     width, height = letter
    
#     # Set font for Hindi text
#     c.setFont(font_name, 12)
    
#     # Split text into lines and write to PDF
#     y_position = height - 50
#     line_height = 14
    
#     # Simple text wrapping
#     words = text.split()
#     current_line = ""
    
#     for word in words:
#         test_line = current_line + " " + word if current_line else word
        
#         # If line would be too long, write the current line and start a new one
#         if c.stringWidth(test_line, font_name, 12) > (width - 100):
#             c.drawString(50, y_position, current_line)
#             y_position -= line_height
#             current_line = word
            
#             # New page if needed
#             if y_position < 50:
#                 c.showPage()
#                 c.setFont(font_name, 12)
#                 y_position = height - 50
#         else:
#             current_line = test_line
    
#     # Write the last line
#     if current_line:
#         c.drawString(50, y_position, current_line)
    
#     c.save()
#     return buffer.getvalue()

# # Function to save text as DOCX
# def save_as_docx(text):
#     doc = Document()
#     doc.add_paragraph(text)
    
#     # Save the document to a bytes buffer
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer.getvalue()

# # Function to create a download link
# def get_download_link(file_bytes, file_name, file_format):
#     b64 = base64.b64encode(file_bytes).decode()
#     mime_types = {
#         'pdf': 'application/pdf',
#         'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
#         'txt': 'text/plain'
#     }
#     mime_type = mime_types.get(file_format, 'application/octet-stream')
#     href = f'<a href="data:{mime_type};base64,{b64}" download="{file_name}">Download {file_name}</a>'
#     return href

# # Main function
# def main():
#     # File uploader
#     uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'txt', 'docx'])
    
#     if uploaded_file is not None:
#         # Display file details
#         file_details = {
#             "Filename": uploaded_file.name,
#             "File size": f"{uploaded_file.size / 1024:.2f} KB",
#             "File type": uploaded_file.type
#         }
#         st.write("### File Details")
#         for key, value in file_details.items():
#             st.write(f"**{key}:** {value}")
        
#         # Extract text based on file type
#         extracted_text = ""
#         file_extension = Path(uploaded_file.name).suffix.lower()
        
#         with st.spinner("Extracting text from document..."):
#             if file_extension == '.pdf':
#                 extracted_text = extract_text_from_pdf(uploaded_file)
#             elif file_extension == '.txt':
#                 extracted_text = uploaded_file.getvalue().decode('utf-8')
#             elif file_extension == '.docx':
#                 extracted_text = extract_text_from_docx(uploaded_file)
        
#         # Display extracted text
#         with st.expander("Show Original Text", expanded=True):
#             st.text_area("Original Text", value=extracted_text, height=200)
        
#         # Translation button
#         if st.button("Translate Document"):
#             with st.spinner("Translating document..."):
#                 translated_text = translate_text(extracted_text)
#                 st.session_state.translated_text = translated_text
#                 st.session_state.file_name = uploaded_file.name
#                 st.session_state.file_extension = file_extension
        
#         # Display translation if available
#         if 'translated_text' in st.session_state:
#             st.markdown("### Translation Result")
            
#             with st.expander("Show Hindi Translation", expanded=True):
#                 st.text_area("Translated Text", value=st.session_state.translated_text, height=200)
            
#             # Download options
#             st.markdown("### Download Options")
            
#             col1, col2, col3 = st.columns(3)
            
#             # Get base filename without extension
#             base_name = Path(st.session_state.file_name).stem
            
#             # PDF download
#             with col1:
#                 if st.button("Download as PDF"):
#                     with st.spinner("Preparing PDF..."):
#                         pdf_bytes = save_as_pdf(st.session_state.translated_text)
#                         st.markdown(get_download_link(pdf_bytes, f"{base_name}_hindi.pdf", "pdf"), unsafe_allow_html=True)
            
#             # DOCX download
#             with col2:
#                 if st.button("Download as DOCX"):
#                     with st.spinner("Preparing DOCX..."):
#                         docx_bytes = save_as_docx(st.session_state.translated_text)
#                         st.markdown(get_download_link(docx_bytes, f"{base_name}_hindi.docx", "docx"), unsafe_allow_html=True)
            
#             # TXT download
#             with col3:
#                 if st.button("Download as TXT"):
#                     with st.spinner("Preparing TXT..."):
#                         txt_bytes = st.session_state.translated_text.encode('utf-8')
#                         st.markdown(get_download_link(txt_bytes, f"{base_name}_hindi.txt", "txt"), unsafe_allow_html=True)
    
#     # Display information about the service
#     st.sidebar.markdown("## About")
#     st.sidebar.info("""
#     This app translates legal documents from English to Hindi.
    
#     **Supported file formats:**
#     - PDF (.pdf)
#     - Text files (.txt)
#     - Word documents (.docx)
    
#     The translation is powered by a specialized legal translation model.
#     """)
    
#     # Add footer
#     st.markdown("---")
#     st.markdown("© 2025 Legal Document Translation Service")

# if __name__ == "__main__":
#     main()


import streamlit as st
import requests
import json
import io
import os
import tempfile
import base64
from pathlib import Path

# Document processing libraries
import PyPDF2
import docx
import fitz  # PyMuPDF
from docx import Document

# Set page config
st.set_page_config(
    page_title="Legal Document Translator (English → Hindi)",
    page_icon="🔄",
    layout="wide"
)

# Application title and description
st.title("Legal Document Translator")
st.markdown("### English to Hindi Legal Document Translation")
st.markdown("Upload a document (.pdf, .txt, or .docx) to translate legal content from English to Hindi.")

# API endpoint
TRANSLATION_API_URL = "https://d89e-2401-4900-8843-57cc-cdb6-af8-25c0-bb20.ngrok-free.app/translate"

# Function to extract text from PDF
def extract_text_from_pdf(file):
    text = ""
    try:
        # Try using PyMuPDF first (better for complex PDFs)
        pdf_document = fitz.open(stream=file.read(), filetype="pdf")
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text += page.get_text()
        pdf_document.close()
    except Exception as e:
        st.warning(f"PyMuPDF extraction failed, trying PyPDF2: {e}")
        # Reset file pointer
        file.seek(0)
        # Fallback to PyPDF2
        pdf_reader = PyPDF2.PdfReader(file)
        for page_num in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page_num].extract_text() or ""
    
    return text

# Function to extract text from DOCX
def extract_text_from_docx(file):
    doc = Document(file)
    return " ".join([paragraph.text for paragraph in doc.paragraphs])

# Function to translate text
def translate_text(text):
    if not text.strip():
        return "No text content found to translate."
    
    try:
        response = requests.post(
            TRANSLATION_API_URL,
            json={"text": text},
            headers={"Content-Type": "application/json"}
        )
        
        if response.status_code == 200:
            return response.json().get("translation", "Translation not found in response")
        else:
            return f"Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"Translation API error: {str(e)}"

# Function to save text as DOCX
def save_as_docx(text):
    doc = Document()
    
    # Split text by paragraphs for better formatting
    paragraphs = text.split('\n\n')
    for para in paragraphs:
        doc.add_paragraph(para)
    
    # Save the document to a bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Function to create a download link
def get_download_link(file_bytes, file_name, file_format):
    b64 = base64.b64encode(file_bytes).decode()
    mime_types = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'txt': 'text/plain'
    }
    mime_type = mime_types.get(file_format, 'application/octet-stream')
    href = f'<a href="data:{mime_type};base64,{b64}" download="{file_name}">Download {file_name}</a>'
    return href

# Main function
def main():
    # File uploader
    uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'txt', 'docx'])
    
    if uploaded_file is not None:
        # Display file details
        file_details = {
            "Filename": uploaded_file.name,
            "File size": f"{uploaded_file.size / 1024:.2f} KB",
            "File type": uploaded_file.type
        }
        st.write("### File Details")
        for key, value in file_details.items():
            st.write(f"**{key}:** {value}")
        
        # Extract text based on file type
        extracted_text = ""
        file_extension = Path(uploaded_file.name).suffix.lower()
        
        with st.spinner("Extracting text from document..."):
            if file_extension == '.pdf':
                extracted_text = extract_text_from_pdf(uploaded_file)
            elif file_extension == '.txt':
                extracted_text = uploaded_file.getvalue().decode('utf-8')
            elif file_extension == '.docx':
                extracted_text = extract_text_from_docx(uploaded_file)
        
        # Display extracted text
        with st.expander("Show Original Text", expanded=True):
            st.text_area("Original Text", value=extracted_text, height=200)
        
        # Translation button
        if st.button("Translate Document"):
            with st.spinner("Translating document..."):
                translated_text = translate_text(extracted_text)
                st.session_state.translated_text = translated_text
                st.session_state.file_name = uploaded_file.name
                st.session_state.file_extension = file_extension
        
        # Display translation if available
        if 'translated_text' in st.session_state:
            st.markdown("### Translation Result")
            
            with st.expander("Show Hindi Translation", expanded=True):
                st.text_area("Translated Text", value=st.session_state.translated_text, height=200)
            
            # Download options
            st.markdown("### Download Options")
            
            col1, col2 = st.columns(2)
            
            # Get base filename without extension
            base_name = Path(st.session_state.file_name).stem
            
            # DOCX download
            with col1:
                if st.button("Download as DOCX"):
                    with st.spinner("Preparing DOCX..."):
                        docx_bytes = save_as_docx(st.session_state.translated_text)
                        st.markdown(get_download_link(docx_bytes, f"{base_name}_hindi.docx", "docx"), unsafe_allow_html=True)
            
            # TXT download
            with col2:
                if st.button("Download as TXT"):
                    with st.spinner("Preparing TXT..."):
                        txt_bytes = st.session_state.translated_text.encode('utf-8')
                        st.markdown(get_download_link(txt_bytes, f"{base_name}_hindi.txt", "txt"), unsafe_allow_html=True)
    
    # Display information about the service
    st.sidebar.markdown("## About")
    st.sidebar.info("""
    This app translates legal documents from English to Hindi.
    
    **Supported file formats:**
    - PDF (.pdf)
    - Text files (.txt)
    - Word documents (.docx)
    
    The translation is powered by a specialized legal translation model.
    """)
    
    # Add footer
    st.markdown("---")
    st.markdown("© 2025 Legal Document Translation Service")

if __name__ == "__main__":
    main()